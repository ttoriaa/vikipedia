#!/usr/bin/env python3
from __future__ import annotations

import argparse
import datetime as dt
import html
import json
import os
import re
import shutil
import subprocess
import textwrap
import time
import zipfile
from pathlib import Path
from typing import Any

import requests

ROOT = Path(__file__).resolve().parents[1]
OUT_ROOT = ROOT / "reports" / "resume_tailoring"

VERSIONS = ["conservative", "balanced", "aggressive"]

ZH_STOPWORDS = {
    "负责", "进行", "以及", "相关", "通过", "包括", "提升", "支持", "推动", "参与", "能力", "经验", "工作", "岗位", "要求",
    "我们", "你将", "能够", "以上", "优先", "熟悉", "具有", "团队", "公司", "业务", "项目", "技术", "数据", "平台",
}

EN_STOPWORDS = {
    "the", "and", "for", "with", "from", "that", "this", "will", "have", "your", "you", "our", "are", "is", "to", "of",
    "in", "on", "or", "as", "be", "an", "a", "by", "at", "we", "job", "role", "work", "team", "years", "year", "experience",
}


def _resolve_llm_api_key() -> str:
    return os.getenv("GLM_API_KEY", "").strip() or os.getenv("OPENAI_API_KEY", "").strip()


def _resolve_llm_base_url(default: str = "https://api.openai.com/v1") -> str:
    configured = os.getenv("OPENAI_BASE_URL", "").strip()
    if configured:
        return configured.rstrip("/")
    if os.getenv("GLM_API_KEY", "").strip():
        return os.getenv("GLM_BASE_URL", "https://open.bigmodel.cn/api/paas/v4").strip().rstrip("/")
    return default.rstrip("/")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Tailor resume content to a JD and export multi-version outputs (HTML/PDF/Word)."
    )
    parser.add_argument("--resume-file", default="", help="Path to resume text file")
    parser.add_argument("--resume-text", default="", help="Raw resume text")
    parser.add_argument("--jd-file", default="", help="Path to JD text file")
    parser.add_argument("--jd-text", default="", help="Raw JD text")
    parser.add_argument("--language", choices=["zh", "en", "bilingual"], default="zh")
    parser.add_argument("--versions", default=",".join(VERSIONS), help="Comma-separated versions")
    parser.add_argument("--model", default="gpt-4o-mini", help="OpenAI model name")
    parser.add_argument("--formats", default="html,pdf,word", help="Output formats: html,pdf,word,json,md")
    parser.add_argument("--out-dir", default="", help="Output directory. Default: reports/resume_tailoring/<date>")
    parser.add_argument("--job-title", default="", help="Optional target job title")
    parser.add_argument(
        "--preserve-structure",
        action="store_true",
        help="Keep original section order and bullet counts; only rewrite text content.",
    )
    parser.add_argument(
        "--resume-template-docx",
        default="",
        help="Path to original resume DOCX template. If set, script writes a DOCX with identical layout and replaced bullet text.",
    )
    parser.add_argument(
        "--template-version",
        choices=VERSIONS,
        default="balanced",
        help="Which rewritten version to inject into DOCX template when --resume-template-docx is used.",
    )
    parser.add_argument(
        "--template-output-name",
        default="resume_tailoring_preserved.docx",
        help="Output filename for layout-preserved DOCX generated from --resume-template-docx.",
    )
    parser.add_argument(
        "--resume-template-html",
        default="",
        help="Path to original resume HTML template. If set, script writes HTML with unchanged layout and replaced list text.",
    )
    parser.add_argument(
        "--html-template-version",
        choices=VERSIONS,
        default="balanced",
        help="Which rewritten version to inject into HTML template when --resume-template-html is used.",
    )
    parser.add_argument(
        "--html-output-name",
        default="resume_tailoring_preserved.html",
        help="Output filename for layout-preserved HTML generated from --resume-template-html.",
    )
    parser.add_argument("--dry-run", action="store_true", help="Generate payload and print summary only")
    return parser.parse_args()


def _clean(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _read_text(path: str) -> str:
    return Path(path).read_text(encoding="utf-8")


def _extract_json_object(raw: str) -> dict[str, Any]:
    text = _clean(raw)
    if not text:
        raise ValueError("Empty model response")

    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise ValueError("No JSON object found")
    parsed = json.loads(match.group(0))
    if not isinstance(parsed, dict):
        raise ValueError("Model response JSON is not an object")
    return parsed


def _html_to_text(html_text: str) -> str:
    text = re.sub(r"<style[\\s\\S]*?</style>", " ", html_text, flags=re.IGNORECASE)
    text = re.sub(r"<script[\\s\\S]*?</script>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html.unescape(text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _load_input(
    resume_file: str,
    resume_text: str,
    jd_file: str,
    jd_text: str,
    resume_template_html: str = "",
) -> tuple[str, str]:
    resume = _clean(resume_text)
    jd = _clean(jd_text)
    if not resume and resume_file:
        resume = _read_text(resume_file)
    if not resume and resume_template_html:
        html_path = Path(resume_template_html)
        if html_path.exists():
            resume = _html_to_text(html_path.read_text(encoding="utf-8"))
    if not jd and jd_file:
        jd = _read_text(jd_file)
    if not resume:
        raise ValueError("Resume text is required via --resume-text, --resume-file, or --resume-template-html")
    if not jd:
        raise ValueError("JD text is required via --jd-text or --jd-file")
    return resume, jd


def _split_lines(text: str) -> list[str]:
    normalized = text
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    # Force common bullet and section separators to new lines.
    normalized = re.sub(r"\s*[•·]\s*", "\n", normalized)
    normalized = re.sub(r"\s{2,}", "\n", normalized)
    normalized = re.sub(r"(?<=。)\s*", "\n", normalized)

    lines = []
    for raw in normalized.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if line:
            lines.append(line)
    return lines


def _is_noise_line(line: str) -> bool:
    lowered = line.lower()
    if any(token in lowered for token in ["phone:", "email:", "@", "gpa", "教育经历", "奖励", "发表论文"]):
        return True
    if re.fullmatch(r"[\d/\-–至今年月\s\.]+", line):
        return True
    if len(line) <= 6:
        return True
    return False


def _clean_bullet_line(line: str, max_len: int = 140) -> str:
    text = re.sub(r"^[\-\*•\d\.\)\(\s]+", "", line).strip(" ;，。")
    text = re.sub(r"\s+", " ", text)
    text = text.replace(" ,", ",")
    if len(text) > max_len:
        text = text[:max_len].rstrip(" ,，。")
    return text


def _extract_original_resume_bullets(resume_text: str) -> dict[str, list[str]]:
    sections = {"work": [], "internship": []}
    state = ""
    heading_stop_tokens = ["教育经历", "发表论文", "奖励", "技能", "projects", "project experience"]

    for raw in resume_text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = _clean(raw)
        if not line:
            continue

        lower = line.lower()
        if "工作经历" in line or "work experience" in lower:
            state = "work"
            continue
        if "实习经历" in line or "internship" in lower:
            state = "internship"
            continue
        if any(token in line or token in lower for token in heading_stop_tokens):
            state = ""
            continue

        if state not in sections:
            continue

        if re.match(r"^[•\-\*]\s*", line):
            sections[state].append(_clean_bullet_line(re.sub(r"^[•\-\*]\s*", "", line), max_len=220))
            continue

        if sections[state] and not _is_noise_line(line):
            sections[state][-1] = _clean_bullet_line(sections[state][-1] + " " + line, max_len=220)

    sections["work"] = [item for item in sections["work"] if item]
    sections["internship"] = [item for item in sections["internship"] if item]
    return sections


def _fit_list_to_count(items: list[str], count: int) -> list[str]:
    if count <= 0:
        return items
    if len(items) >= count:
        return items[:count]
    if not items:
        return [""] * count
    out = list(items)
    idx = 0
    while len(out) < count:
        out.append(items[idx % len(items)])
        idx += 1
    return out[:count]


def _extract_docx_template_bullets(template_docx: Path) -> dict[str, list[dict[str, Any]]]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("python-docx is required for --resume-template-docx. Install via: pip install python-docx") from exc

    doc = Document(str(template_docx))
    result: dict[str, list[dict[str, Any]]] = {"work": [], "internship": []}
    state = ""
    stop_tokens = ["教育经历", "发表论文", "奖励", "技能", "projects", "project experience"]

    for idx, paragraph in enumerate(doc.paragraphs):
        text = _clean(paragraph.text)
        if not text:
            continue

        lower = text.lower()
        if "工作经历" in text or "work experience" in lower:
            state = "work"
            continue
        if "实习经历" in text or "internship" in lower:
            state = "internship"
            continue
        if any(token in text or token in lower for token in stop_tokens):
            state = ""
            continue

        if state not in result:
            continue

        if re.match(r"^[•\-\*]\s*", text):
            raw = re.sub(r"^[•\-\*]\s*", "", text)
            result[state].append(
                {
                    "index": idx,
                    "raw_text": text,
                    "clean_text": _clean_bullet_line(raw, max_len=220),
                }
            )

    return result


def _replace_paragraph_text_preserve_layout(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.text = text


def _write_docx_with_preserved_layout(
    template_docx: Path,
    output_docx: Path,
    payload: dict[str, Any],
    version: str,
) -> tuple[bool, str]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        return False, f"python-docx is required for --resume-template-docx: {exc}"

    if version not in payload.get("versions", {}):
        return False, f"Template version not found in payload: {version}"

    template_map = _extract_docx_template_bullets(template_docx)
    work_template = template_map.get("work", [])
    intern_template = template_map.get("internship", [])

    version_block = payload["versions"][version]
    work_rewrite = [
        _clean_bullet_line(item, max_len=220)
        for item in version_block.get("work_experience", [])
        if _clean(item)
    ]
    intern_rewrite = [
        _clean_bullet_line(item, max_len=220)
        for item in version_block.get("internship_experience", [])
        if _clean(item)
    ]

    work_rewrite = _fit_list_to_count(work_rewrite, len(work_template))
    intern_rewrite = _fit_list_to_count(intern_rewrite, len(intern_template))

    doc = Document(str(template_docx))

    for pos, item in enumerate(work_template):
        idx = item["index"]
        original = item.get("clean_text", "")
        rewritten = _clean(work_rewrite[pos]) if pos < len(work_rewrite) else ""
        final_text = rewritten or original
        bullet_prefix = "• "
        _replace_paragraph_text_preserve_layout(doc.paragraphs[idx], bullet_prefix + final_text)

    for pos, item in enumerate(intern_template):
        idx = item["index"]
        original = item.get("clean_text", "")
        rewritten = _clean(intern_rewrite[pos]) if pos < len(intern_rewrite) else ""
        final_text = rewritten or original
        bullet_prefix = "• "
        _replace_paragraph_text_preserve_layout(doc.paragraphs[idx], bullet_prefix + final_text)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return True, "ok"


def _extract_list_items_from_ul(ul_block: str) -> list[str]:
    items: list[str] = []
    for match in re.finditer(r"<li[^>]*>([\s\S]*?)</li>", ul_block, flags=re.IGNORECASE):
        raw = match.group(1)
        txt = _clean_bullet_line(_html_to_text(raw), max_len=220)
        if txt:
            items.append(txt)
    return items


def _replace_section_ul_lists(section_html: str, list_values: list[list[str]]) -> str:
    idx = 0

    def _repl(match: re.Match[str]) -> str:
        nonlocal idx
        if idx >= len(list_values):
            idx += 1
            return match.group(0)
        values = list_values[idx]
        idx += 1
        li_lines = "\n".join(f"            <li>{html.escape(item)}</li>" for item in values)
        return "<ul>\n" + li_lines + "\n          </ul>"

    return re.sub(r"<ul>[\s\S]*?</ul>", _repl, section_html, flags=re.IGNORECASE)


def _write_html_with_preserved_layout(
    template_html_path: Path,
    output_html_path: Path,
    payload: dict[str, Any],
    version: str,
) -> tuple[bool, str]:
    if version not in payload.get("versions", {}):
        return False, f"Template version not found in payload: {version}"

    html_text = template_html_path.read_text(encoding="utf-8")

    work_match = re.search(
        r"(<section[^>]*>[\s\S]*?<h2>\s*工作经历\s*</h2>[\s\S]*?</section>)",
        html_text,
        flags=re.IGNORECASE,
    )
    intern_match = re.search(
        r"(<section[^>]*>[\s\S]*?<h2>\s*实习经历\s*</h2>[\s\S]*?</section>)",
        html_text,
        flags=re.IGNORECASE,
    )

    if not work_match and not intern_match:
        return False, "No 工作经历 / 实习经历 section found in HTML template"

    block = payload["versions"][version]
    work_rewrite = [
        _clean_bullet_line(item, max_len=220)
        for item in block.get("work_experience", [])
        if _clean(item)
    ]
    intern_rewrite = [
        _clean_bullet_line(item, max_len=220)
        for item in block.get("internship_experience", [])
        if _clean(item)
    ]

    updated_html = html_text

    if work_match:
        section = work_match.group(1)
        ul_blocks = re.findall(r"<ul>[\s\S]*?</ul>", section, flags=re.IGNORECASE)
        original_lists = [_extract_list_items_from_ul(ul) for ul in ul_blocks]
        original_flat = [item for items in original_lists for item in items]
        total = sum(len(items) for items in original_lists)
        source_items = work_rewrite if work_rewrite else original_flat
        work_fitted = _fit_list_to_count(source_items, total)

        split_lists: list[list[str]] = []
        pos = 0
        for items in original_lists:
            count = len(items)
            split_lists.append(work_fitted[pos:pos + count])
            pos += count

        replaced = _replace_section_ul_lists(section, split_lists)
        updated_html = updated_html.replace(section, replaced, 1)

    if intern_match:
        section = intern_match.group(1)
        ul_blocks = re.findall(r"<ul>[\s\S]*?</ul>", section, flags=re.IGNORECASE)
        original_lists = [_extract_list_items_from_ul(ul) for ul in ul_blocks]
        original_flat = [item for items in original_lists for item in items]
        total = sum(len(items) for items in original_lists)
        source_items = intern_rewrite if intern_rewrite else original_flat
        intern_fitted = _fit_list_to_count(source_items, total)

        split_lists = []
        pos = 0
        for items in original_lists:
            count = len(items)
            split_lists.append(intern_fitted[pos:pos + count])
            pos += count

        replaced = _replace_section_ul_lists(section, split_lists)
        updated_html = updated_html.replace(section, replaced, 1)

    output_html_path.parent.mkdir(parents=True, exist_ok=True)
    output_html_path.write_text(updated_html, encoding="utf-8")
    return True, "ok"


def _find_section(lines: list[str], names: list[str]) -> int:
    lowered = [line.lower() for line in lines]
    for idx, line in enumerate(lowered):
        for name in names:
            if name.lower() in line:
                return idx
    return -1


def _extract_resume_sections(resume_text: str) -> dict[str, list[str]]:
    lines = _split_lines(resume_text)
    if not lines:
        return {"work": [], "project": [], "full": []}

    work_start = _find_section(lines, ["工作经历", "work experience", "professional experience"])
    project_start = _find_section(lines, ["项目经历", "projects", "project experience"])
    intern_start = _find_section(lines, ["实习经历", "internship", "internships"])
    award_start = _find_section(lines, ["奖励", "awards", "publications", "技能", "skills"])

    if work_start == -1:
        work_start = 0
    work_end_candidates = [idx for idx in [project_start, intern_start, award_start] if idx != -1 and idx > work_start]
    work_end = min(work_end_candidates) if work_end_candidates else len(lines)

    project_lines: list[str] = []
    if project_start != -1:
        project_end_candidates = [idx for idx in [intern_start, award_start] if idx != -1 and idx > project_start]
        project_end = min(project_end_candidates) if project_end_candidates else len(lines)
        project_lines = lines[project_start + 1 : project_end]

    work_lines = lines[work_start + 1 : work_end]

    # Remove obvious noise and preserve meaningful content lines only.
    work_lines = [line for line in work_lines if not _is_noise_line(line)]
    project_lines = [line for line in project_lines if not _is_noise_line(line)]
    full_lines = [line for line in lines if not _is_noise_line(line)]

    return {"work": work_lines, "project": project_lines, "full": full_lines}


def _tokenize_keywords(text: str) -> list[str]:
    zh_tokens = re.findall(r"[\u4e00-\u9fff]{2,8}", text)
    en_tokens = re.findall(r"\b[A-Za-z][A-Za-z0-9+/\-]{2,}\b", text)

    out: list[str] = []
    for token in zh_tokens:
        if token not in ZH_STOPWORDS:
            out.append(token)
    for token in en_tokens:
        t = token.lower()
        if t not in EN_STOPWORDS:
            out.append(token)
    return out


def _extract_top_keywords(jd_text: str, top_n: int = 30) -> list[str]:
    freq: dict[str, int] = {}
    for token in _tokenize_keywords(jd_text):
        freq[token] = freq.get(token, 0) + 1
    ranked = sorted(freq.items(), key=lambda item: (-item[1], -len(item[0]), item[0]))
    return [token for token, _ in ranked[:top_n]]


def _line_score(line: str, keywords: list[str]) -> int:
    text = line.lower()
    score = 0
    for kw in keywords:
        if kw.lower() in text:
            score += 2
    if re.search(r"\d", line):
        score += 1
    if any(ch in line for ch in ["主导", "牵头", "lead", "owned", "delivered"]):
        score += 2
    return score


def _pick_candidate_bullets(lines: list[str], keywords: list[str], max_items: int = 10) -> list[str]:
    cleaned: list[str] = []
    for line in lines:
        if _is_noise_line(line):
            continue
        line = _clean_bullet_line(line)
        if len(line) < 10:
            continue
        if any(token in line.lower() for token in ["graduate trainee program", "technology risk consultant"]):
            continue
        if line:
            cleaned.append(line)

    ranked = sorted(cleaned, key=lambda item: (_line_score(item, keywords), len(item)), reverse=True)
    seen: set[str] = set()
    picked: list[str] = []
    for line in ranked:
        key = line[:80]
        if key in seen:
            continue
        seen.add(key)
        picked.append(line)
        if len(picked) >= max_items:
            break
    return picked


def _rewrite_bullet(line: str, style: str, language: str) -> str:
    base = _clean_bullet_line(line)
    base = re.sub(r"^(围绕|主导)+", "", base).strip(" ，。")
    if style == "conservative":
        return base

    if language == "en":
        if style == "balanced":
            templates = [
                "Delivered {body} with clear ownership and coordinated execution across stakeholders",
                "Drove {body} with disciplined planning, cross-team alignment, and reliable delivery quality",
                "Executed {body} through structured prioritization and consistent delivery follow-through",
            ]
            idx = sum(ord(ch) for ch in base) % len(templates)
            body = base[0].lower() + base[1:] if len(base) > 1 else base.lower()
            return templates[idx].format(body=body)

        templates = [
            "Led {body} end to end, aligning business priorities with execution milestones and outcomes",
            "Owned {body} from planning to rollout, driving alignment, delivery speed, and measurable impact",
            "Spearheaded {body} with strong stakeholder orchestration and execution rigor tied to business goals",
            "Directed {body} as a full-cycle initiative, balancing strategic goals with dependable execution",
        ]
        idx = sum(ord(ch) for ch in base) % len(templates)
        body = base[0].lower() + base[1:] if len(base) > 1 else base.lower()
        return templates[idx].format(body=body)

    if style == "balanced":
        templates = [
            "{body}，并通过结构化推进与跨团队协作确保交付稳定",
            "{body}，结合节奏化推进与协同机制提升落地效率",
            "{body}，在明确目标与里程碑的基础上实现高质量交付",
        ]
        idx = sum(ord(ch) for ch in base) % len(templates)
        return templates[idx].format(body=base)

    templates = [
        "主导{body}的端到端落地，拉通跨团队资源并确保关键里程碑兑现",
        "围绕{body}建立执行闭环，兼顾业务目标、交付效率与质量稳定性",
        "牵引{body}从方案到上线的全流程推进，在复杂协作中持续放大业务结果",
        "统筹{body}的落地节奏与协同机制，以结果导向推动项目高质量完成",
    ]
    idx = sum(ord(ch) for ch in base) % len(templates)
    return templates[idx].format(body=base)


def _fallback_payload(resume_text: str, jd_text: str, language: str, versions: list[str], job_title: str) -> dict[str, Any]:
    sections = _extract_resume_sections(resume_text)
    keywords = _extract_top_keywords(jd_text)
    work_candidates = _pick_candidate_bullets(sections["work"] or sections["full"], keywords, max_items=9)
    project_candidates = _pick_candidate_bullets(sections["project"], keywords, max_items=5)

    if not project_candidates:
        derived = [line for line in work_candidates if any(k in line for k in ["项目", "平台", "交付", "需求", "协同"])]
        project_candidates = derived[:4]
    if not project_candidates:
        project_candidates = [
            "组织跨团队需求澄清、方案设计与落地验证，形成可复用的项目推进路径",
            "建立里程碑管理与风险跟踪机制，提升项目透明度与交付稳定性",
        ]

    focus_items = keywords[:10]
    version_map: dict[str, Any] = {}

    for version in versions:
        version_map[version] = {
            "positioning": (
                f"{job_title}岗位匹配版本" if job_title else "面向目标JD的匹配版本"
            ),
            "work_experience": [_rewrite_bullet(line, version, language) for line in work_candidates[:6]],
            "project_experience": [_rewrite_bullet(line, version, language) for line in project_candidates[:4]],
            "internship_experience": [],
            "keyword_coverage": focus_items[:8],
        }

    return {
        "meta": {
            "provider": "fallback-rule",
            "language": language,
            "job_title": job_title,
        },
        "jd_focus": focus_items,
        "gaps": ["建议补充可量化结果（如效率、成本、周期、质量指标）以提升说服力。"],
        "versions": version_map,
    }


def _build_prompt(resume_text: str, jd_text: str, language: str, versions: list[str], job_title: str) -> str:
    version_text = ", ".join(versions)
    lang_rule = {
        "zh": "请用中文输出。",
        "en": "Please output in English.",
        "bilingual": "先输出中文再输出英文，对应结构保持一致。",
    }[language]

    return textwrap.dedent(
        f"""
        你是资深简历优化顾问。任务：基于给定简历与JD，产出三种版本改写（{version_text}）。

        关键约束：
        1) 只能使用简历中已有事实，不得虚构公司、职责、项目、数字。
        2) 重点改写工作经历与项目经历，强调与JD匹配的关键词、职责、业务结果。
        3) 句子要可直接粘贴到简历，避免空话和夸张表达。
        4) 每条内容尽量动词开头，信息密度高。
        5) {lang_rule}

        仅返回严格 JSON，结构如下：
        {{
          "jd_focus": ["..."],
          "gaps": ["..."],
          "versions": {{
            "conservative": {{
              "positioning": "...",
              "work_experience": ["..."],
              "project_experience": ["..."],
              "keyword_coverage": ["..."]
            }},
            "balanced": {{...}},
            "aggressive": {{...}}
          }}
        }}

        其中：
        - work_experience 每版 4-7 条
        - project_experience 每版 3-5 条
        - keyword_coverage 每版 6-12 个

        目标岗位：{job_title or "未指定"}

        JD全文：
        {jd_text[:12000]}

        简历全文：
        {resume_text[:14000]}
        """
    ).strip()


def _generate_with_openai(
    resume_text: str,
    jd_text: str,
    language: str,
    versions: list[str],
    model: str,
    job_title: str,
) -> dict[str, Any]:
    api_key = _resolve_llm_api_key()
    if not api_key:
        return _fallback_payload(resume_text, jd_text, language, versions, job_title)

    base_url = _resolve_llm_base_url("https://api.openai.com/v1")
    endpoint = f"{base_url}/chat/completions"

    payload = {
        "model": model,
        "temperature": 0.3,
        "messages": [
            {"role": "system", "content": "You are a precise resume tailoring assistant. Output JSON only."},
            {
                "role": "user",
                "content": _build_prompt(resume_text, jd_text, language, versions, job_title),
            },
        ],
        "response_format": {"type": "json_object"},
    }

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    try:
        response = requests.post(endpoint, headers=headers, data=json.dumps(payload, ensure_ascii=False), timeout=180)
        response.raise_for_status()
        body = response.json()
        content = _clean(body.get("choices", [{}])[0].get("message", {}).get("content"))
        parsed = _extract_json_object(content)
        if "versions" not in parsed:
            raise ValueError("Model output missing versions")
        parsed["meta"] = {
            "provider": "openai",
            "model": model,
            "language": language,
            "job_title": job_title,
        }
        return parsed
    except Exception as exc:  # noqa: BLE001
        data = _fallback_payload(resume_text, jd_text, language, versions, job_title)
        data["meta"]["fallback_reason"] = str(exc)
        return data


def _normalize_payload(payload: dict[str, Any], versions: list[str]) -> dict[str, Any]:
    result = {
        "meta": payload.get("meta", {}),
        "jd_focus": payload.get("jd_focus", []),
        "gaps": payload.get("gaps", []),
        "versions": {},
    }
    for version in versions:
        block = payload.get("versions", {}).get(version, {})
        result["versions"][version] = {
            "positioning": _clean(block.get("positioning")) or "",
            "work_experience": [
                _clean(item) for item in block.get("work_experience", []) if _clean(item)
            ],
            "project_experience": [
                _clean(item) for item in block.get("project_experience", []) if _clean(item)
            ],
            "internship_experience": [
                _clean(item) for item in block.get("internship_experience", []) if _clean(item)
            ],
            "keyword_coverage": [
                _clean(item) for item in block.get("keyword_coverage", []) if _clean(item)
            ],
        }
    return result


def _apply_preserve_structure(
    payload: dict[str, Any],
    resume_text: str,
    versions: list[str],
    language: str,
) -> dict[str, Any]:
    original = _extract_original_resume_bullets(resume_text)
    work_original = original.get("work", [])
    internship_original = original.get("internship", [])

    for version in versions:
        work_rewrite = [_rewrite_bullet(item, version, language) for item in work_original] if work_original else []
        intern_rewrite = [_rewrite_bullet(item, version, language) for item in internship_original] if internship_original else []

        payload["versions"][version]["work_experience"] = _fit_list_to_count(work_rewrite, len(work_original))
        payload["versions"][version]["internship_experience"] = _fit_list_to_count(intern_rewrite, len(internship_original))

    payload["preserve_structure"] = {
        "work_bullet_count": len(work_original),
        "internship_bullet_count": len(internship_original),
    }
    return payload


def _render_markdown(payload: dict[str, Any], versions: list[str], job_title: str) -> str:
    lines = [
        "# Resume Tailoring Output",
        "",
        f"- Generated at: {dt.datetime.now().isoformat(timespec='seconds')}",
        f"- Job title: {job_title or 'N/A'}",
        f"- Provider: {payload.get('meta', {}).get('provider', 'unknown')}",
        "",
        "## JD Focus",
    ]
    for item in payload.get("jd_focus", []):
        lines.append(f"- {item}")

    lines.extend(["", "## Gap Notes"])
    for item in payload.get("gaps", []):
        lines.append(f"- {item}")

    for version in versions:
        block = payload["versions"][version]
        lines.extend(["", f"## Version: {version}"])
        if block["positioning"]:
            lines.append(f"- Positioning: {block['positioning']}")

        lines.extend(["", "### Work Experience"])
        for item in block["work_experience"]:
            lines.append(f"- {item}")

        lines.extend(["", "### Project Experience"])
        for item in block["project_experience"]:
            lines.append(f"- {item}")

        if block.get("internship_experience"):
            lines.extend(["", "### Internship Experience"])
            for item in block["internship_experience"]:
                lines.append(f"- {item}")

        lines.extend(["", "### Keyword Coverage"])
        for item in block["keyword_coverage"]:
            lines.append(f"- {item}")

    return "\n".join(lines).strip()


def _render_html(payload: dict[str, Any], versions: list[str], job_title: str) -> str:
    focus_html = "".join(f"<li>{html.escape(item)}</li>" for item in payload.get("jd_focus", []))
    gap_html = "".join(f"<li>{html.escape(item)}</li>" for item in payload.get("gaps", []))

    cards = []
    for version in versions:
        block = payload["versions"][version]
        work_items = "".join(f"<li>{html.escape(item)}</li>" for item in block["work_experience"])
        project_items = "".join(f"<li>{html.escape(item)}</li>" for item in block["project_experience"])
        intern_items = "".join(f"<li>{html.escape(item)}</li>" for item in block.get("internship_experience", []))
        keywords = "".join(f"<span class=\"chip\">{html.escape(item)}</span>" for item in block["keyword_coverage"])

        cards.append(
            f"""
            <section class=\"card\">
              <h2>{html.escape(version.title())}</h2>
              <p class=\"muted\">{html.escape(block['positioning'])}</p>
              <h3>Work Experience</h3>
              <ul>{work_items}</ul>
              <h3>Project Experience</h3>
              <ul>{project_items}</ul>
                            <h3>Internship Experience</h3>
                            <ul>{intern_items}</ul>
              <h3>Keyword Coverage</h3>
              <div class=\"chips\">{keywords}</div>
            </section>
            """
        )

    generated = dt.datetime.now().isoformat(timespec="seconds")
    provider = html.escape(_clean(payload.get("meta", {}).get("provider", "unknown")))

    return f"""<!doctype html>
<html lang=\"en\">
<head>
  <meta charset=\"UTF-8\" />
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\" />
  <title>Resume Tailoring Output</title>
  <style>
    :root {{
      --bg: #f4f3ef;
      --ink: #1f2a33;
      --muted: #5f6a72;
      --card: #ffffff;
      --line: #d6d8db;
      --accent: #0a5a7a;
      --chip: #e7f1f5;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      color: var(--ink);
      font-family: "Segoe UI", "Trebuchet MS", Arial, sans-serif;
      background: radial-gradient(circle at 8% 6%, rgba(10, 90, 122, 0.14), transparent 45%), var(--bg);
      padding: 18px;
    }}
    .wrap {{ max-width: 1120px; margin: 0 auto; display: grid; gap: 14px; }}
    .hero {{ background: linear-gradient(126deg, #123b52, #1b4e4f); color: #f7fbfc; border-radius: 18px; padding: 20px; }}
    .hero p {{ margin: 6px 0; }}
    .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); gap: 12px; }}
    .card {{ background: var(--card); border: 1px solid var(--line); border-radius: 14px; padding: 16px; }}
    h1 {{ margin: 0; font-size: 30px; }}
    h2 {{ margin: 0 0 6px; font-size: 22px; color: var(--accent); }}
    h3 {{ margin: 12px 0 6px; font-size: 15px; }}
    p, li {{ line-height: 1.55; }}
    .muted {{ color: var(--muted); }}
    ul {{ margin: 6px 0 0; padding-left: 18px; }}
    .chips {{ display: flex; flex-wrap: wrap; gap: 6px; margin-top: 6px; }}
    .chip {{ background: var(--chip); border: 1px solid #cde3ec; border-radius: 999px; padding: 4px 10px; font-size: 12px; }}
    @media print {{
      body {{ padding: 0; background: #fff; }}
      .hero {{ border-radius: 0; }}
      .card {{ break-inside: avoid; }}
    }}
  </style>
</head>
<body>
  <main class=\"wrap\">
    <section class=\"hero\">
      <h1>Resume Tailoring Output</h1>
      <p>Job title: {html.escape(job_title or 'N/A')}</p>
      <p>Generated at: {generated}</p>
      <p>Provider: {provider}</p>
    </section>

    <section class=\"card\">
      <h2>JD Focus</h2>
      <ul>{focus_html}</ul>
      <h3>Gap Notes</h3>
      <ul>{gap_html}</ul>
    </section>

    <div class=\"grid\">
      {''.join(cards)}
    </div>
  </main>
</body>
</html>
"""


def _docx_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _make_docx(payload: dict[str, Any], versions: list[str], output_path: Path, job_title: str) -> None:
    paragraphs: list[str] = []
    paragraphs.append("Resume Tailoring Output")
    paragraphs.append(f"Job title: {job_title or 'N/A'}")
    paragraphs.append(f"Provider: {payload.get('meta', {}).get('provider', 'unknown')}")
    paragraphs.append("")
    paragraphs.append("JD Focus")
    paragraphs.extend(f"- {item}" for item in payload.get("jd_focus", []))
    paragraphs.append("")
    paragraphs.append("Gap Notes")
    paragraphs.extend(f"- {item}" for item in payload.get("gaps", []))

    for version in versions:
        block = payload["versions"][version]
        paragraphs.append("")
        paragraphs.append(f"Version: {version}")
        if block["positioning"]:
            paragraphs.append(f"Positioning: {block['positioning']}")
        paragraphs.append("Work Experience")
        paragraphs.extend(f"- {item}" for item in block["work_experience"])
        paragraphs.append("Project Experience")
        paragraphs.extend(f"- {item}" for item in block["project_experience"])
        if block.get("internship_experience"):
            paragraphs.append("Internship Experience")
            paragraphs.extend(f"- {item}" for item in block["internship_experience"])
        paragraphs.append("Keyword Coverage")
        paragraphs.extend(f"- {item}" for item in block["keyword_coverage"])

    body = []
    for text in paragraphs:
        safe = _docx_escape(text)
        if not safe:
            body.append("<w:p/>")
            continue
        body.append(
            "<w:p><w:r><w:t xml:space=\"preserve\">"
            + safe
            + "</w:t></w:r></w:p>"
        )

    document_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document xmlns:wpc=\"http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas\" "
        "xmlns:mc=\"http://schemas.openxmlformats.org/markup-compatibility/2006\" "
        "xmlns:o=\"urn:schemas-microsoft-com:office:office\" "
        "xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\" "
        "xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\" "
        "xmlns:v=\"urn:schemas-microsoft-com:vml\" "
        "xmlns:wp14=\"http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing\" "
        "xmlns:wp=\"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing\" "
        "xmlns:w10=\"urn:schemas-microsoft-com:office:word\" "
        "xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
        "xmlns:w14=\"http://schemas.microsoft.com/office/word/2010/wordml\" "
        "xmlns:w15=\"http://schemas.microsoft.com/office/word/2012/wordml\" "
        "xmlns:wpg=\"http://schemas.microsoft.com/office/word/2010/wordprocessingGroup\" "
        "xmlns:wpi=\"http://schemas.microsoft.com/office/word/2010/wordprocessingInk\" "
        "xmlns:wne=\"http://schemas.microsoft.com/office/word/2006/wordml\" "
        "xmlns:wps=\"http://schemas.microsoft.com/office/word/2010/wordprocessingShape\" "
        "mc:Ignorable=\"w14 w15 wp14\">"
        "<w:body>"
        + "".join(body)
        + "<w:sectPr><w:pgSz w:w=\"11906\" w:h=\"16838\"/>"
        "<w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\" w:header=\"720\" w:footer=\"720\" w:gutter=\"0\"/>"
        "</w:sectPr></w:body></w:document>"
    )

    content_types_xml = """<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>
<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">
  <Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>
  <Default Extension=\"xml\" ContentType=\"application/xml\"/>
  <Override PartName=\"/word/document.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>
</Types>
"""

    rels_xml = """<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>
<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">
  <Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/>
</Relationships>
"""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types_xml)
        zf.writestr("_rels/.rels", rels_xml)
        zf.writestr("word/document.xml", document_xml)


def _find_browser_exe() -> str | None:
    candidates = [
        shutil.which("msedge"),
        shutil.which("msedge.exe"),
        shutil.which("chrome"),
        shutil.which("chrome.exe"),
        shutil.which("chromium"),
        shutil.which("chromium.exe"),
        os.path.join(os.getenv("ProgramFiles", ""), "Microsoft", "Edge", "Application", "msedge.exe"),
        os.path.join(os.getenv("ProgramFiles(x86)", ""), "Microsoft", "Edge", "Application", "msedge.exe"),
        os.path.join(os.getenv("ProgramFiles", ""), "Google", "Chrome", "Application", "chrome.exe"),
        os.path.join(os.getenv("ProgramFiles(x86)", ""), "Google", "Chrome", "Application", "chrome.exe"),
    ]
    for item in candidates:
        if item and Path(item).exists():
            return item
    return None


def _export_pdf_from_html(html_path: Path, pdf_path: Path) -> tuple[bool, str]:
    browser = _find_browser_exe()
    if not browser:
        return False, "No Edge/Chrome executable found for headless PDF export"

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    cmd = [
        browser,
        "--headless=new",
        "--disable-gpu",
        "--no-first-run",
        "--no-default-browser-check",
        "--allow-file-access-from-files",
        f"--print-to-pdf={str(pdf_path.resolve())}",
        html_path.resolve().as_uri(),
    ]
    try:
        proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        timeout_s = 90
        start = time.monotonic()

        while True:
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                # Some Edge builds keep running even after file is written.
                proc.terminate()
                try:
                    proc.wait(timeout=5)
                except subprocess.TimeoutExpired:
                    proc.kill()
                return True, "ok"

            if proc.poll() is not None:
                break

            if time.monotonic() - start > timeout_s:
                proc.terminate()
                try:
                    proc.wait(timeout=5)
                except subprocess.TimeoutExpired:
                    proc.kill()
                if pdf_path.exists() and pdf_path.stat().st_size > 0:
                    return True, "ok"
                return False, "Browser timed out before producing PDF"

            time.sleep(0.5)

        stdout_text, stderr_text = proc.communicate(timeout=2)
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True, "ok"
        return False, f"Browser PDF export failed: {(stderr_text or stdout_text).strip()}"
    except Exception as exc:  # noqa: BLE001
        return False, str(exc)


def _write_outputs(payload: dict[str, Any], versions: list[str], out_dir: Path, formats: set[str], job_title: str) -> dict[str, Any]:
    out_dir.mkdir(parents=True, exist_ok=True)
    results: dict[str, Any] = {"files": {}, "warnings": []}

    if "json" in formats:
        path = out_dir / "resume_tailoring.json"
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        results["files"]["json"] = str(path)

    md = _render_markdown(payload, versions, job_title)
    if "md" in formats:
        path = out_dir / "resume_tailoring.md"
        path.write_text(md + "\n", encoding="utf-8")
        results["files"]["md"] = str(path)

    html_text = _render_html(payload, versions, job_title)
    html_path = out_dir / "resume_tailoring.html"
    if "html" in formats or "pdf" in formats:
        html_path.write_text(html_text + "\n", encoding="utf-8")
        results["files"]["html"] = str(html_path)

    if "word" in formats:
        docx_path = out_dir / "resume_tailoring.docx"
        _make_docx(payload, versions, docx_path, job_title)
        results["files"]["word"] = str(docx_path)

    if "pdf" in formats:
        pdf_path = out_dir / "resume_tailoring.pdf"
        ok, message = _export_pdf_from_html(html_path, pdf_path)
        if ok:
            results["files"]["pdf"] = str(pdf_path)
        else:
            results["warnings"].append(message)

    return results


def main() -> int:
    args = parse_args()
    resume_text, jd_text = _load_input(
        args.resume_file,
        args.resume_text,
        args.jd_file,
        args.jd_text,
        args.resume_template_html,
    )

    versions = [item.strip() for item in args.versions.split(",") if item.strip()]
    versions = [item for item in versions if item in VERSIONS] or VERSIONS

    formats = {item.strip().lower() for item in args.formats.split(",") if item.strip()}
    formats = formats.intersection({"html", "pdf", "word", "json", "md"})
    if not formats:
        formats = {"html", "pdf", "word"}

    run_date = dt.date.today().isoformat()
    out_dir = Path(args.out_dir) if args.out_dir else (OUT_ROOT / run_date)

    raw_payload = _generate_with_openai(
        resume_text=resume_text,
        jd_text=jd_text,
        language=args.language,
        versions=versions,
        model=args.model,
        job_title=args.job_title,
    )
    payload = _normalize_payload(raw_payload, versions)
    # In HTML template mode, structure is preserved by replacing list content in-place,
    # so we should not remap bullets from plain-text extraction.
    if args.preserve_structure and not args.resume_template_html:
        payload = _apply_preserve_structure(payload, resume_text, versions, args.language)

    if args.dry_run:
        print(f"Provider: {payload.get('meta', {}).get('provider', 'unknown')}")
        print(f"Versions: {', '.join(versions)}")
        print(f"JD focus count: {len(payload.get('jd_focus', []))}")
        for version in versions:
            block = payload["versions"][version]
            print(
                f"{version}: work={len(block['work_experience'])}, "
                f"internship={len(block.get('internship_experience', []))}, "
                f"project={len(block['project_experience'])}, "
                f"keywords={len(block['keyword_coverage'])}"
            )
        return 0

    output_result = _write_outputs(payload, versions, out_dir, formats, args.job_title)

    if args.resume_template_docx:
        template_path = Path(args.resume_template_docx)
        if not template_path.exists():
            output_result.setdefault("warnings", []).append(
                f"Template DOCX not found: {template_path}"
            )
        else:
            target_name = _clean(args.template_output_name) or "resume_tailoring_preserved.docx"
            if not target_name.lower().endswith(".docx"):
                target_name = target_name + ".docx"
            preserved_docx = out_dir / target_name
            ok, msg = _write_docx_with_preserved_layout(
                template_docx=template_path,
                output_docx=preserved_docx,
                payload=payload,
                version=args.template_version,
            )
            if ok:
                output_result.setdefault("files", {})["word_preserved"] = str(preserved_docx)
            else:
                output_result.setdefault("warnings", []).append(msg)

    if args.resume_template_html:
        template_html = Path(args.resume_template_html)
        if not template_html.exists():
            output_result.setdefault("warnings", []).append(
                f"Template HTML not found: {template_html}"
            )
        else:
            target_name = _clean(args.html_output_name) or "resume_tailoring_preserved.html"
            if not target_name.lower().endswith(".html"):
                target_name = target_name + ".html"
            preserved_html = out_dir / target_name
            ok, msg = _write_html_with_preserved_layout(
                template_html_path=template_html,
                output_html_path=preserved_html,
                payload=payload,
                version=args.html_template_version,
            )
            if ok:
                output_result.setdefault("files", {})["html_preserved"] = str(preserved_html)
            else:
                output_result.setdefault("warnings", []).append(msg)

    print(f"Provider: {payload.get('meta', {}).get('provider', 'unknown')}")
    print(f"Output directory: {out_dir}")
    for name, path in output_result.get("files", {}).items():
        print(f"{name}: {path}")
    for warn in output_result.get("warnings", []):
        print(f"warning: {warn}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
